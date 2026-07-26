"""
parser.py
Handles extraction of raw text from uploaded resume files (PDF / DOCX).
"""

import os
import fitz  # PyMuPDF
import docx  # python-docx


def extract_text_from_pdf(file_path):
    """Extract raw text from a PDF file using PyMuPDF."""
    text = ""
    try:
        pdf_document = fitz.open(file_path)
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            text += page.get_text("text") + "\n"
        pdf_document.close()
    except Exception as exc:
        raise ValueError(f"Failed to read PDF file: {exc}")
    return text.strip()


def extract_text_from_docx(file_path):
    """Extract raw text from a DOCX file using python-docx."""
    text = ""
    try:
        document = docx.Document(file_path)
        for para in document.paragraphs:
            text += para.text + "\n"

        # Also extract text from tables (many resumes use table layouts)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as exc:
        raise ValueError(f"Failed to read DOCX file: {exc}")
    return text.strip()


def extract_text(file_path):
    """
    Detect file type by extension and extract text accordingly.
    Returns the extracted plain text.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")


def allowed_file(filename, allowed_extensions={"pdf", "docx"}):
    """Check whether the uploaded filename has an allowed extension."""
    return (
        "." in filename
        and filename.rsplit(".", 1)[1].lower() in allowed_extensions
    )
